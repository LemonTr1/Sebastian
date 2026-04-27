from docx import Document
import typer
from agents import function_tool

@function_tool
async def modify_docx(
        file_path: str,
        new_content: str,
        mode: str = "append",
        old_content: str = None
) -> str:
    """
    修改 docx 文档的内容。
    Args:
        file_path: docx 文件的完整路径
        new_content: 要添加或替换的内容
        mode: 操作模式 —— "append"（追加到末尾）、"replace"（全局替换文本）
        old_content: replace 模式下需要替换的目标文本（必填）
    Returns:
        操作结果说明（字符串）
    """
    typer.echo(typer.style("[执行中] 正在修改 docx 文档内容", fg=typer.colors.WHITE))

    try:
        doc = Document(file_path)

        if mode == "append":
            doc.add_paragraph(new_content)

        elif mode == "replace":
            if old_content is None:
                return "错误：replace 模式必须提供 old_content 参数，用于指定要被替换的内容"

            replaced_count = 0
            for paragraph in doc.paragraphs:
                if old_content in paragraph.text:
                    # 保留样式：逐段重建 runs
                    full_text = paragraph.text.replace(old_content, new_content)

                    # 保存原段落对齐方式等格式
                    paragraph_alignment = paragraph.paragraph_format.alignment

                    # 清空 runs
                    paragraph.clear()

                    # 重建文本
                    run = paragraph.add_run(full_text)

                    # 恢复段落格式
                    paragraph.paragraph_format.alignment = paragraph_alignment

                    replaced_count += 1

            if replaced_count == 0:
                return f"警告：未在文档中找到文本 '{old_content}'，未执行任何替换"

        else:
            return f"错误：不支持的 mode: {mode}，仅支持 'append' 或 'replace'"

        # 交互确认
        confirmed = typer.confirm(
            typer.style(f"[Warn] 是否要保存对 {file_path} 的修改？", fg=typer.colors.YELLOW)
        )
        if not confirmed:
            typer.echo("操作已终止")
            return "操作已取消，文件未保存"

        doc.save(file_path)

        if mode == "append":
            return f"成功修改文档：已追加内容到末尾"
        else:
            return f"成功修改文档：replace 模式已完成，共替换 {replaced_count} 处"

    except FileNotFoundError:
        return f"错误：文件不存在 - {file_path}"
    except PermissionError:
        return f"错误：文件被占用或无写入权限 - {file_path}"
    except Exception as e:
        return f"修改文档时出错：{str(e)}"