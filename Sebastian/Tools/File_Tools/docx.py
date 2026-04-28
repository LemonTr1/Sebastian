from docx import Document
from agents import function_tool
import typer
import asyncio

#创建空docx文档
@function_tool
async def create_docx(file_path: str) -> str:
    """
    在指定路径创建一个全新的空白 .docx 文档。
    该文档不包含任何预设内容（无标题、无段落），
    可被 LibreOffice 和 Microsoft Office 正常打开。
    Args:
        file_path: 要创建的 .docx 文件的完整路径
    Returns:
        操作结果说明（成功或失败）
    """
    try:
        confirmed = typer.confirm(typer.style(f"[Warn]确定创建/覆盖文档{file_path}吗？",fg=typer.colors.YELLOW))
        if not confirmed:
            typer.echo("已终止此操作")
            return f"用户阻止了本次操作"
        # 创建一个空白的 Document 对象
        doc = Document()
        # 直接保存，不添加任何内容
        doc.save(file_path)
        typer.echo(typer.style(f"{file_path}已创建", fg=typer.colors.WHITE))
        return f"成功创建空白文档：{file_path}"
    except Exception as e:
        return f"创建失败：{str(e)}"

#读取docx文档内容
@function_tool
async def read_docx(file_path: str) -> str:
    """
    读取docx文档的纯文本内容。
    Args:
        file_path: docx文件的完整路径
    Returns:
        成功时返回：文档的纯文本内容
        错误时返回：报错信息
    """
    try:
        typer.echo(typer.style(f"[执行中]正在读取{file_path}文档内容...",fg=typer.colors.WHITE))
        loop = asyncio.get_running_loop()
        doc = await loop.run_in_executor(None, Document, file_path)
        # 收集所有段落的文本
        paragraphs_text = []
        for para in doc.paragraphs:
            if para.text.strip():  # 过滤空段落
                paragraphs_text.append(para.text)
        # 也可以收集表格内容（可选），python-docx 可以方便地遍历表格并将其内容一并返回[reference:0]
        tables_text = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append("\t".join(row_data))
            if table_data:
                tables_text.append("\n".join(table_data))

        # 合并段落和表格内容返回
        full_text = "\n\n".join(paragraphs_text)
        if tables_text:
            full_text += "\n\n【表格内容】\n" + "\n\n".join(tables_text)

        return full_text if full_text else "文档中没有可读取的文本内容。"
    except FileNotFoundError:
        return f"错误：文件不存在 - {file_path}"
    except Exception as e:
        return f"读取文档时出错：{str(e)}"

#修改docx文档
@function_tool
async def modify_docx(
        file_path: str,
        new_content: str,
        mode: str = "append",
        old_content: str = None
) -> str:
    """
    修改 docx 文档的内容。
    Args:
        file_path: docx 文件的完整路径
        new_content: 要添加或替换的内容
        mode: 操作模式 —— "append"（追加到末尾）、"replace"（全局替换文本）
        old_content: replace 模式下需要替换的目标文本（必填）
    Returns:
        操作结果说明（字符串）
    """
    typer.echo(typer.style(f"[执行中] 正在修改{file_path}文档内容", fg=typer.colors.WHITE))

    try:
        doc = Document(file_path)

        if mode == "append":
            doc.add_paragraph(new_content)

        elif mode == "replace":
            if old_content is None:
                return "错误：replace 模式必须提供 old_content 参数，用于指定要被替换的内容"

            replaced_count = 0
            for paragraph in doc.paragraphs:
                if old_content in paragraph.text:
                    # 保留样式：逐段重建 runs
                    full_text = paragraph.text.replace(old_content, new_content)

                    # 保存原段落对齐方式等格式
                    paragraph_alignment = paragraph.paragraph_format.alignment

                    # 清空 runs
                    paragraph.clear()

                    # 重建文本
                    run = paragraph.add_run(full_text)

                    # 恢复段落格式
                    paragraph.paragraph_format.alignment = paragraph_alignment

                    replaced_count += 1

            if replaced_count == 0:
                return f"警告：未在文档中找到文本 '{old_content}'，未执行任何替换"

        else:
            return f"错误：不支持的 mode: {mode}，仅支持 'append' 或 'replace'"

        # 交互确认
        confirmed = typer.confirm(
            typer.style(f"[Warn] 是否要保存对 {file_path} 的修改？", fg=typer.colors.YELLOW)
        )
        if not confirmed:
            typer.echo("操作已终止")
            return "操作已取消，文件未保存"

        doc.save(file_path)

        if mode == "append":
            return f"成功修改文档：已追加内容到末尾"
        else:
            return f"成功修改文档：replace 模式已完成，共替换 {replaced_count} 处"

    except FileNotFoundError:
        return f"错误：文件不存在 - {file_path}"
    except PermissionError:
        return f"错误：文件被占用或无写入权限 - {file_path}"
    except Exception as e:
        return f"修改文档时出错：{str(e)}"