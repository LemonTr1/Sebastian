import os
import json
from docx import Document
from docx.shared import Pt
from src.security.path_safety import resolve_safe_path
from src.utils.exceptions import SecurityException
from src.tools.tools_registry import get_tools_registry


def _docx_to_text(doc) -> str:
    lines = []
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            lines.append(f"## {para.text}")
        else:
            lines.append(para.text)
    for i, table in enumerate(doc.tables):
        lines.append(f"\n[表格 {i + 1}]")
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells]
            lines.append(" | ".join(row_cells))
    return "\n".join(lines)


def _parse_content(content: str) -> list:
    try:
        parsed = json.loads(content)
    except json.JSONDecodeError as e:
        raise ValueError(f"content 参数 JSON 解析失败：{e}")
    if not isinstance(parsed, list):
        raise ValueError("content 必须是 JSON 数组格式")
    return parsed


def _populate_document(doc, blocks: list):
    for block in blocks:
        if not isinstance(block, dict):
            continue
        block_type = block.get("type", "paragraph")
        text = block.get("text", "")
        if block_type == "title":
            doc.add_heading(text, level=0)
        elif block_type == "heading":
            level = int(block.get("level", 1))
            level = max(1, min(level, 9))
            doc.add_heading(text, level=level)
        elif block_type == "paragraph":
            para = doc.add_paragraph()
            run = para.add_run(text)
            if block.get("bold"):
                run.bold = True
            if block.get("italic"):
                run.italic = True
            if block.get("underline"):
                run.underline = True
            size_pt = block.get("font_size")
            if size_pt:
                run.font.size = Pt(int(size_pt))
        elif block_type == "table":
            rows_data = block.get("rows", [])
            if not rows_data:
                continue
            max_cols = max(len(r) for r in rows_data) if rows_data else 1
            table = doc.add_table(rows=len(rows_data), cols=max_cols)
            table.style = 'Table Grid'
            for r_idx, row_data in enumerate(rows_data):
                for c_idx, cell_text in enumerate(row_data):
                    cell = table.cell(r_idx, c_idx)
                    cell.text = str(cell_text)


def read_docx(path: str, filename: str) -> str:
    if not filename.lower().endswith(".docx"):
        return json.dumps(
            {
                "success": False,
                "summary": "仅支持 .docx 文件",
                "content": None
            },
            ensure_ascii=False
        )
    path = os.path.abspath(path)
    file_path = os.path.join(path, filename)
    try:
        file_path = resolve_safe_path(file_path)
    except SecurityException as e:
        return json.dumps(
            {
                "success": False,
                "summary": str(e),
                "content": None
            },
            ensure_ascii=False
        )
    try:
        doc = Document(file_path)
        content = _docx_to_text(doc)
        return json.dumps(
            {
                "success": True,
                "summary": f"成功读取 {file_path}",
                "content": content
            },
            ensure_ascii=False
        )
    except Exception as e:
        return json.dumps(
            {
                "success": False,
                "summary": str(e),
                "content": None
            },
            ensure_ascii=False
        )


def create_docx(path: str, filename: str, content: str) -> str:
    if not filename.lower().endswith(".docx"):
        return json.dumps(
            {
                "success": False,
                "summary": "文件名须以 .docx 结尾"
            },
            ensure_ascii=False
        )
    path = os.path.abspath(path)
    try:
        path = resolve_safe_path(path)
    except SecurityException as e:
        return json.dumps(
            {
                "success": False,
                "summary": str(e)
            },
            ensure_ascii=False
        )
    if not os.path.isdir(path):
        return json.dumps(
            {
                "success": False,
                "summary": f"父目录不存在: {path}"
            },
            ensure_ascii=False
        )
    file_path = os.path.join(path, filename)
    if os.path.exists(file_path):
        return json.dumps(
            {
                "success": False,
                "summary": f"文件已存在: {file_path}，请用 write_docx 覆盖"
            },
            ensure_ascii=False
        )
    try:
        blocks = _parse_content(content)
    except ValueError as e:
        return json.dumps(
            {
                "success": False,
                "summary": str(e)
            },
            ensure_ascii=False
        )
    try:
        doc = Document()
        _populate_document(doc, blocks)
        doc.save(file_path)
        return json.dumps(
            {
                "success": True,
                "summary": f"docx 创建成功: {file_path}"
            },
            ensure_ascii=False
        )
    except Exception as e:
        return json.dumps(
            {
                "success": False,
                "summary": str(e)
            },
            ensure_ascii=False
        )


def write_docx(path: str, filename: str, content: str) -> str:
    if not filename.lower().endswith(".docx"):
        return json.dumps(
            {
                "success": False,
                "summary": "文件名须以 .docx 结尾"
            },
            ensure_ascii=False
        )
    path = os.path.abspath(path)
    try:
        path = resolve_safe_path(path)
    except SecurityException as e:
        return json.dumps(
            {
                "success": False,
                "summary": str(e)
            },
            ensure_ascii=False
        )
    if not os.path.isdir(path):
        return json.dumps(
            {
                "success": False,
                "summary": f"父目录不存在: {path}"
            },
            ensure_ascii=False
        )
    file_path = os.path.join(path, filename)
    try:
        blocks = _parse_content(content)
    except ValueError as e:
        return json.dumps(
            {
                "success": False,
                "summary": str(e)
            },
            ensure_ascii=False
        )
    existed = os.path.exists(file_path)
    try:
        doc = Document()
        _populate_document(doc, blocks)
        doc.save(file_path)
        return json.dumps(
            {
                "success": True,
                "summary": f"docx {'覆盖' if existed else '创建'}成功: {file_path}"
            },
            ensure_ascii=False
        )
    except Exception as e:
        return json.dumps(
            {
                "success": False,
                "summary": str(e)
            },
            ensure_ascii=False
        )


def edit_docx(path: str, filename: str, content: str) -> str:
    if not filename.lower().endswith(".docx"):
        return json.dumps(
            {
                "success": False,
                "summary": "仅支持 .docx 文件"
            },
            ensure_ascii=False
        )
    path = os.path.abspath(path)
    file_path = os.path.join(path, filename)
    try:
        file_path = resolve_safe_path(file_path)
    except SecurityException as e:
        return json.dumps(
            {
                "success": False,
                "summary": str(e)
            },
            ensure_ascii=False
        )
    if not os.path.isfile(file_path):
        return json.dumps(
            {
                "success": False,
                "summary": f"文件不存在: {file_path}"
            },
            ensure_ascii=False
        )
    try:
        operations = _parse_content(content)
    except ValueError as e:
        return json.dumps(
            {
                "success": False,
                "summary": str(e)
            },
            ensure_ascii=False
        )
    performed = []
    try:
        doc = Document(file_path)
        for op in operations:
            action = op.get("action")
            if action == "find_replace":
                find_text = op.get("find", "")
                replace_text = op.get("replace", "")
                if not find_text:
                    continue
                replaced_count = 0
                for para in doc.paragraphs:
                    if find_text not in para.text:
                        continue
                    full = ""
                    for run in para.runs:
                        full += run.text
                    if find_text in full:
                        new_full = full.replace(find_text, replace_text)
                        for run in para.runs:
                            run.text = ""
                        if para.runs:
                            para.runs[0].text = new_full
                        replaced_count += 1
                performed.append(f"替换 '{find_text}'->'{replace_text}' 命中 {replaced_count} 处")
            elif action == "append":
                blocks = op.get("blocks", [])
                if blocks:
                    _populate_document(doc, blocks)
                    performed.append(f"追加 {len(blocks)} 个内容块")
            elif action == "replace_all":
                blocks = op.get("blocks", [])
                if blocks:
                    for para in list(doc.paragraphs):
                        p = para._element
                        p.getparent().remove(p)
                    for table in list(doc.tables):
                        tbl = table._element
                        tbl.getparent().remove(tbl)
                    _populate_document(doc, blocks)
                    performed.append(f"整体替换为 {len(blocks)} 个内容块")
            else:
                performed.append(f"跳过未知操作: {action}")
        doc.save(file_path)
        return json.dumps(
            {
                "success": True,
                "summary": "; ".join(performed) if performed else "未执行操作"
            },
            ensure_ascii=False
        )
    except Exception as e:
        return json.dumps(
            {
                "success": False,
                "summary": str(e)
            },
            ensure_ascii=False
        )

READ_DOCX_SCHEMA = {
    "type": "function",
    "function": {
        "name": "read_docx",
        "description": "读取 .docx 文件的全部文本内容（段落+表格）",
        "parameters": {
            "type": "object",
            "properties": {
                "path": {"type": "string", "description": "父目录绝对路径"},
                "filename": {"type": "string", "description": ".docx 文件名"},
            },
            "required": ["path", "filename"],
        },
    },
}

CREATE_DOCX_SCHEMA = {
    "type": "function",
    "function": {
        "name": "create_docx",
        "description": "创建新 .docx 文件（目标不能已存在）。content 为 JSON 数组：[{\"type\":\"title\",\"text\":\"...\"},{\"type\":\"heading\",\"level\":2,\"text\":\"...\"},{\"type\":\"paragraph\",\"text\":\"...\",\"bold\":false,\"italic\":false},{\"type\":\"table\",\"rows\":[[\"H1\",\"H2\"],[\"C1\",\"C2\"]]}]",
        "parameters": {
            "type": "object",
            "properties": {
                "path": {"type": "string", "description": "父目录绝对路径"},
                "filename": {"type": "string", "description": ".docx 文件名"},
                "content": {"type": "string", "description": "JSON 数组字符串，定义文档结构"},
            },
            "required": ["path", "filename", "content"],
        },
    },
}

WRITE_DOCX_SCHEMA = {
    "type": "function",
    "function": {
        "name": "write_docx",
        "description": "创建或覆盖写入 .docx 文件（存在时直接覆盖）。content 格式同 create_docx。【此工具需要用户确认后方可执行——覆盖已有文件不可逆】",
        "parameters": {
            "type": "object",
            "properties": {
                "path": {"type": "string", "description": "父目录绝对路径"},
                "filename": {"type": "string", "description": ".docx 文件名"},
                "content": {"type": "string", "description": "JSON 数组字符串"},
            },
            "required": ["path", "filename", "content"],
        },
    },
}

EDIT_DOCX_SCHEMA = {
    "type": "function",
    "function": {
        "name": "edit_docx",
        "description": "编辑已有 .docx 文件。content 为编辑操作JSON数组，支持: find_replace(查找替换), append(尾部追加), replace_all(整体替换)。【此工具需要用户确认后方可执行】",
        "parameters": {
            "type": "object",
            "properties": {
                "path": {"type": "string", "description": "父目录绝对路径"},
                "filename": {"type": "string", "description": ".docx 文件名"},
                "content": {"type": "string", "description": "JSON 数组字符串，每项为编辑操作"},
            },
            "required": ["path", "filename", "content"],
        },
    },
}

get_tools_registry().register_tool("read_docx", read_docx, READ_DOCX_SCHEMA, for_agent="File_Agent")
get_tools_registry().register_tool("create_docx", create_docx, CREATE_DOCX_SCHEMA, for_agent="File_Agent")
get_tools_registry().register_tool("write_docx", write_docx, WRITE_DOCX_SCHEMA, hitl=True, for_agent="File_Agent")
get_tools_registry().register_tool("edit_docx", edit_docx, EDIT_DOCX_SCHEMA, hitl=True, for_agent="File_Agent")
